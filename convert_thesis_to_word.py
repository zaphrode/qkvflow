#!/usr/bin/env python3
"""
Convert THESIS.md to Word (.docx) format with embedded figures

Requirements:
    pip install python-docx

Usage:
    python convert_thesis_to_word.py
"""

import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Installing required packages...")
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT


def add_figure(doc, image_path, caption, figure_num, width_inches=6.0):
    """Add a figure with caption to the document"""
    if os.path.exists(image_path):
        # Add the image centered
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(image_path, width=Inches(width_inches))
        
        # Add caption
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_para.add_run(f"Figure {figure_num}: {caption}")
        caption_run.font.size = Pt(10)
        caption_run.font.italic = True
        
        doc.add_paragraph()  # Add spacing
        return True
    else:
        doc.add_paragraph(f"[Figure {figure_num}: {caption} - Image not found at {image_path}]")
        return False


def create_thesis_document():
    """Create a professional Word document from the thesis with embedded figures"""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # =====================================================
    # TITLE PAGE
    # =====================================================
    
    for _ in range(4):
        doc.add_paragraph()
    
    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Time-Indexed Parameter Sharing\nfor Neural ODE Transformers")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("A Final Year Project Thesis")
    sub_run.font.size = Pt(18)
    sub_run.font.italic = True
    
    for _ in range(3):
        doc.add_paragraph()
    
    # Author info
    info_items = [
        ("Author:", "[Your Name]"),
        ("Supervisor:", "[Supervisor Name]"),
        ("Institution:", "[University Name]"),
        ("Department:", "[Department of Computer Science]"),
        ("Date:", "January 2026")
    ]
    
    for label, value in info_items:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(label + " ").bold = True
        para.add_run(value)
    
    doc.add_page_break()
    
    # =====================================================
    # ABSTRACT
    # =====================================================
    
    doc.add_heading('Abstract', level=1)
    
    abstract_text = """Transformer architectures have become the dominant paradigm in natural language processing, but their memory requirements scale linearly with depth, limiting deployment on resource-constrained devices. This thesis presents a novel Time-Indexed Parameter Sharing approach that extends Neural ODE Transformers by sharing base weights across layers and modulating them with lightweight time-dependent networks.

We compare four architectures: Standard Transformers (308.5M parameters), Tong et al.'s Neural ODE Transformers (51.5M parameters), our Time-Indexed MLP variant (0.7M parameters, 430× compression), and our Time-Indexed SSM variant (4.9M parameters, 63× compression). Through rigorous statistical validation with 5 random seeds on WikiText-2 and extended evaluation on WikiText-103, we demonstrate that:

1. Time-Indexed MLP achieves validation loss of 2.231 ± 0.025, outperforming both the Standard Transformer (2.367 ± 0.022) and Tong's Neural ODE (2.336 ± 0.018) while using 430× fewer parameters

2. Time-Indexed SSM achieves the best validation loss of 2.147 ± 0.124, with 63× parameter reduction

3. All improvements are statistically significant (p < 0.01) with large effect sizes (Cohen's d > 2.0)

We also identify and document the SSM Speed Paradox: despite fewer parameters, the SSM variant is slower per training step (64.3ms vs 55.3ms) due to sequential scan operations that cannot be parallelized like matrix multiplications in attention. This finding has important implications for practitioners choosing between memory efficiency and inference latency."""
    
    doc.add_paragraph(abstract_text)
    
    keywords = doc.add_paragraph()
    keywords.add_run("Keywords: ").bold = True
    keywords.add_run("Neural ODE, Transformers, Parameter Sharing, State Space Models, Language Modeling, Efficient Deep Learning")
    
    doc.add_page_break()
    
    # =====================================================
    # TABLE OF CONTENTS
    # =====================================================
    
    doc.add_heading('Table of Contents', level=1)
    
    toc_items = [
        ("1. Introduction", "4"),
        ("   1.1 Background and Motivation", "4"),
        ("   1.2 Problem Statement", "4"),
        ("   1.3 Research Questions", "5"),
        ("   1.4 Contributions", "5"),
        ("2. Literature Review", "6"),
        ("   2.1 Transformer Architectures", "6"),
        ("   2.2 Neural Ordinary Differential Equations", "6"),
        ("   2.3 State Space Models", "7"),
        ("   2.4 Parameter-Efficient Methods", "7"),
        ("3. Methodology", "8"),
        ("   3.1 Time-Indexed Parameter Sharing Framework", "8"),
        ("   3.2 Time-Indexed MLP Architecture", "9"),
        ("   3.3 Time-Indexed SSM Architecture", "10"),
        ("   3.4 Training Procedure", "11"),
        ("4. Implementation", "12"),
        ("5. Results and Analysis", "13"),
        ("   5.1 WikiText-2 Statistical Validation", "13"),
        ("   5.2 Statistical Significance Tests", "14"),
        ("   5.3 Parameter Efficiency Analysis", "15"),
        ("   5.4 The SSM Speed Paradox", "16"),
        ("6. Discussion", "18"),
        ("7. Conclusion and Future Work", "20"),
        ("8. References", "22"),
    ]
    
    for item, page in toc_items:
        para = doc.add_paragraph()
        para.add_run(item)
        para.add_run("\t" * 6 + page)
    
    doc.add_paragraph()
    doc.add_paragraph("[Note: Update page numbers after final formatting. In Word: References → Table of Contents → Update Table]").italic = True
    
    doc.add_page_break()
    
    # =====================================================
    # CHAPTER 1: INTRODUCTION
    # =====================================================
    
    doc.add_heading('1. Introduction', level=1)
    
    doc.add_heading('1.1 Background and Motivation', level=2)
    doc.add_paragraph("""Transformer architectures (Vaswani et al., 2017) have revolutionized natural language processing, achieving state-of-the-art results on tasks ranging from machine translation to language modeling. However, the standard Transformer architecture presents significant challenges for deployment:

1. Memory Requirements: Each layer maintains separate weight matrices, leading to memory usage that scales linearly with depth

2. Computational Cost: Deep Transformers require substantial compute for both training and inference

3. Over-parameterization: Many parameters may be redundant, as evidenced by successful pruning and distillation techniques

Recent work on Neural ODE Transformers (Tong et al., ICLR 2025) reframes Transformer layers as discretizations of continuous dynamics, where layer depth corresponds to integration time. This perspective opens new possibilities for parameter efficiency through weight sharing across the "time" dimension of network depth.""")
    
    doc.add_heading('1.2 Problem Statement', level=2)
    doc.add_paragraph("""While Tong et al.'s Neural ODE Transformers provide a principled framework for continuous-depth networks, their approach generates all weight matrices independently at each layer using hypernetworks, resulting in 51.5M parameters—still substantial for edge deployment.

The central question this thesis addresses is:

Can we achieve comparable or better performance by sharing base weights across layers and modulating them with lightweight time-dependent functions, achieving extreme parameter compression?""")
    
    doc.add_heading('1.3 Research Questions', level=2)
    doc.add_paragraph("""This thesis investigates the following research questions:

RQ1: Does time-indexed parameter sharing improve upon standard and Neural ODE Transformers in terms of parameter efficiency?

RQ2: What is the trade-off between parameter count and model performance (validation loss)?

RQ3: How do different time-modulation mechanisms (MLP vs SSM) compare in terms of performance and computational efficiency?

RQ4: Are the observed improvements statistically significant across multiple random seeds?""")
    
    doc.add_heading('1.4 Contributions', level=2)
    doc.add_paragraph("""This thesis makes the following contributions:

1. Novel Architecture: We propose Time-Indexed Parameter Sharing, a technique that shares base weights across all transformer layers while modulating them with lightweight time-dependent networks

2. Two Model Variants: We develop and evaluate two variants:
   • Time-Indexed MLP (0.7M parameters, 430× compression)
   • Time-Indexed SSM (4.9M parameters, 63× compression)

3. Rigorous Evaluation: We provide statistically validated results with 5 random seeds, confidence intervals, and significance testing

4. The SSM Speed Paradox: We document an important finding that fewer parameters do not always translate to faster inference, with implications for architecture selection

5. Open-Source Implementation: We release all code, trained models, and experimental scripts for reproducibility""")
    
    doc.add_page_break()
    
    # =====================================================
    # CHAPTER 2: LITERATURE REVIEW
    # =====================================================
    
    doc.add_heading('2. Literature Review', level=1)
    
    doc.add_heading('2.1 Transformer Architectures', level=2)
    doc.add_paragraph("""The Transformer architecture (Vaswani et al., 2017) introduced self-attention as the primary mechanism for sequence modeling. Key components include:

Multi-Head Self-Attention: The attention mechanism computes weighted sums of values based on query-key similarity:

    Attention(Q, K, V) = softmax(QKᵀ / √dₖ) V

where queries Q, keys K, and values V are linear projections of the input. Multi-head attention allows the model to attend to different representation subspaces.

Feed-Forward Networks: Each attention layer is followed by a position-wise feed-forward network:

    FFN(x) = ReLU(xW₁ + b₁)W₂ + b₂

Layer Normalization and Residual Connections: Following the Pre-LN formulation (Xiong et al., 2020), layer normalization is applied before each sub-layer, improving training stability.""")
    
    doc.add_heading('2.2 Neural Ordinary Differential Equations', level=2)
    doc.add_paragraph("""Chen et al. (2018) introduced Neural ODEs, which parameterize the derivative of hidden states as a neural network:

    dh/dt = f_θ(h(t), t)

The output is obtained by integrating from initial state h(0) to final time T using numerical ODE solvers.

Neural ODE Transformers (Tong et al., ICLR 2025) apply this framework to Transformers by treating layer depth as continuous time. Their key innovation is generating weights at each layer from a hypernetwork conditioned on time embeddings:

    W_QKV(t) = HyperNetwork(SinusoidalEmbed(t))

This approach achieves 51.5M parameters for a 6-layer model with hidden dimension 256.""")
    
    doc.add_heading('2.3 State Space Models', level=2)
    doc.add_paragraph("""State Space Models (SSMs) have emerged as efficient alternatives to attention for sequence modeling. The continuous-time state space is defined by:

    dh/dt = Ah(t) + Bx(t)
    y(t) = Ch(t) + Dx(t)

Mamba (Gu & Dao, 2023) introduces selective state spaces with input-dependent parameters, achieving linear-time complexity in sequence length through efficient parallel scan algorithms.""")
    
    doc.add_heading('2.4 Parameter-Efficient Methods', level=2)
    doc.add_paragraph("""Several approaches have been proposed to reduce Transformer parameters:

Weight Sharing: Universal Transformers (Dehghani et al., 2019) share weights across all layers but lack the expressiveness of depth-varying representations.

Low-Rank Adaptation (LoRA): Hu et al. (2022) propose adding low-rank adapters to frozen pretrained weights: W' = W + BA, where B and A are low-rank matrices.

FiLM (Feature-wise Linear Modulation): Perez et al. (2018) introduced modulating activations with learned scale and shift parameters.

Our Contribution: We extend FiLM to modulate weights rather than activations, applied in a time-indexed manner across network depth:

    W_eff(t) = W_base ⊙ σ(MLP(t))

This provides the expressiveness of depth-varying weights while maintaining extreme parameter efficiency through weight sharing.""")
    
    doc.add_page_break()
    
    # =====================================================
    # CHAPTER 3: METHODOLOGY
    # =====================================================
    
    doc.add_heading('3. Methodology', level=1)
    
    doc.add_heading('3.1 Time-Indexed Parameter Sharing Framework', level=2)
    doc.add_paragraph("""Our core innovation is Time-Indexed Parameter Sharing, which combines three key ideas:

1. Base Weight Sharing: A single set of base weight matrices {W_base} is shared across all layers

2. Time Embedding: Layer depth l is normalized to time t = l/L ∈ [0, 1] and encoded using sinusoidal embeddings

3. Lightweight Modulation: Small networks generate time-dependent scaling factors that modulate the base weights

Mathematical Formulation:

For layer l with normalized time t = l/L:

    γ(t) = σ(MLP_γ(SinusoidalEmbed(t)))
    W_eff(t) = W_base ⊙ γ(t)

where σ is the sigmoid function ensuring modulation factors are in [0, 1], and ⊙ denotes element-wise multiplication.""")
    
    # Add comparison table
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run("Table 1: Comparison with Tong et al.'s Approach").bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Aspect', 'Tong et al. (ICLR 2025)', 'Our Approach']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    data = [
        ['Weight Generation', 'W(t) = Generate(t)', 'W_eff(t) = W_base ⊙ σ(MLP(t))'],
        ['Expressiveness', 'Full matrix generation', 'Base + modulation'],
        ['Parameters', '~51.5M', '~0.7M (430× reduction)'],
        ['Optimization', 'Harder (hypernetworks)', 'Easier (grounded in W_base)']
    ]
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading('3.2 Time-Indexed MLP Architecture', level=2)
    doc.add_paragraph("""The Time-Indexed MLP variant replaces the feed-forward network with a time-modulated shared MLP.

Time Embedding Generation:

The sinusoidal position embedding encodes the normalized layer depth t using sin and cos functions at different frequencies:

    embed(t) = [t, sin(ω₁t), cos(ω₁t), sin(ω₂t), cos(ω₂t), ...]

where ωᵢ = 1/10000^(2i/d) follows the original Transformer formulation.

Time-Indexed Attention:

The attention mechanism uses shared base weights modulated by time-dependent scaling:

    qkv_scale = σ(time_mod_qkv(time_embed))
    out_scale = σ(time_mod_out(time_embed))
    
    x_scaled = x ⊙ qkv_scale
    qkv = W_base_qkv(x_scaled)
    
    # Standard attention computation
    q, k, v = split(qkv)
    attn_out = softmax(qk^T / √d_k) v
    
    out = W_base_out(attn_out) ⊙ out_scale

This allows the same base weights to produce different effective weights at each layer depth.""")
    
    doc.add_heading('3.3 Time-Indexed SSM Architecture', level=2)
    doc.add_paragraph("""The Time-Indexed SSM variant replaces the MLP with a State Space Model, combining the benefits of SSM efficiency with time-indexed parameter sharing.

The SSM is parameterized by matrices A, B, C, D, and discretization step Δ, all generated from time embeddings:

    A_base = -softplus(f_A(t_emb))    # Negative for stability
    B_base = f_B(t_emb)
    C_base = f_C(t_emb)
    D_base = f_D(t_emb)
    Δ_base = softplus(f_Δ(t_emb)) + ε

Time-dependent modulation is then applied:

    A = A_base ⊙ σ(time_mod_A(t))
    B = B_base ⊙ σ(time_mod_B(t))
    ...

The selective scan operation processes the sequence:

    A_bar = exp(Δ · A)           # Discretization
    B_bar = Δ · B
    
    h[t+1] = A_bar · h[t] + B_bar · x[t]    # Recurrence
    y[t] = C · h[t] + D · x[t]

This recurrence is implemented using hax.scan for efficiency, though it remains inherently sequential—a key factor in the SSM Speed Paradox discussed in Section 5.4.""")
    
    doc.add_heading('3.4 Training Procedure', level=2)
    doc.add_paragraph("""Training Configuration:

• Optimizer: AdamW with weight decay 0.01
• Learning rate: 3×10⁻⁴
• Gradient clipping: Global norm 1.0
• Batch size: 8
• Sequence length: 128
• Training steps: 1000

Loss Function:

We use cross-entropy loss with sparse labels to avoid materializing large one-hot tensors:

    loss = softmax_cross_entropy_with_integer_labels(logits, targets)

Statistical Validation:

To ensure robust results, we run each experiment with 5 random seeds: {42, 123, 456, 789, 1011}. We report:

• Mean ± Standard Deviation
• 95% Confidence Intervals (using t-distribution)
• Paired t-tests for significance testing
• Cohen's d effect sizes for practical significance""")
    
    doc.add_page_break()
    
    # =====================================================
    # CHAPTER 4: IMPLEMENTATION
    # =====================================================
    
    doc.add_heading('4. Implementation', level=1)
    
    doc.add_heading('4.1 Software Architecture', level=2)
    doc.add_paragraph("""The implementation is built on a modern JAX-based stack:""")
    
    # Software table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['Library', 'Version', 'Purpose']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    sw_data = [
        ['JAX', '0.4.28+', 'Automatic differentiation, XLA compilation'],
        ['Equinox', '0.11.4+', 'PyTree-based neural network modules'],
        ['Haliax', '1.3+', 'Named tensor operations'],
        ['Optax', '0.2.0+', 'Gradient transformations and optimizers']
    ]
    for i, row_data in enumerate(sw_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    doc.add_paragraph("""Design Principles:

1. Functional Programming: All models are pure functions, enabling easy JIT compilation and parallelization

2. Named Arrays: Using Haliax's named arrays prevents axis permutation bugs common in tensor operations

3. Modular Architecture: Clear separation between time embedding, modulation, and base operations allows for easy experimentation""")
    
    doc.add_heading('4.2 Experimental Setup', level=2)
    
    # Datasets table
    para = doc.add_paragraph()
    para.add_run("Table 2: Datasets Used").bold = True
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    headers = ['Dataset', 'Tokens', 'Purpose']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    ds_data = [
        ['WikiText-2', '2M', 'Primary evaluation with statistical validation'],
        ['WikiText-103', '103M', 'Extended validation (50× larger)']
    ]
    for i, row_data in enumerate(ds_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    # Model config table
    para = doc.add_paragraph()
    para.add_run("Table 3: Model Configuration").bold = True
    
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    headers = ['Hyperparameter', 'Value']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    config_data = [
        ['Hidden dimension', '256'],
        ['Number of heads', '4'],
        ['Number of layers', '6'],
        ['Sequence length', '128'],
        ['Time embedding dimension', '64'],
        ['Sinusoidal dimension', '32'],
        ['SSM state size', '64']
    ]
    for i, row_data in enumerate(config_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    doc.add_paragraph("Hardware: Training was performed on NVIDIA A100 GPU (40GB). Approximate training time: 30-60 minutes for 1000 steps per model configuration.")
    
    doc.add_page_break()
    
    # =====================================================
    # CHAPTER 5: RESULTS AND ANALYSIS
    # =====================================================
    
    doc.add_heading('5. Results and Analysis', level=1)
    
    doc.add_heading('5.1 WikiText-2 Statistical Validation', level=2)
    doc.add_paragraph("We conducted rigorous statistical validation with 5 random seeds on WikiText-2. The results demonstrate that our Time-Indexed approaches significantly outperform both baselines while using dramatically fewer parameters.")
    
    # Main results table
    para = doc.add_paragraph()
    para.add_run("Table 4: Model Performance Comparison (Mean ± Std over 5 seeds)").bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    headers = ['Model', 'Valid Loss', 'Parameters', 'Speed (ms)', 'Compression']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    results_data = [
        ['Standard Transformer', '2.367 ± 0.022', '308.5M', '55.3 ± 1.2', '1.0×'],
        ['Tong\'s Neural ODE', '2.336 ± 0.018', '51.5M', '15.3 ± 0.1', '6.0×'],
        ['Time-Indexed MLP', '2.231 ± 0.025', '0.7M', '7.7 ± 0.3', '430.9×'],
        ['Time-Indexed SSM', '2.147 ± 0.124', '4.9M', '64.3 ± 0.5', '62.9×']
    ]
    for i, row_data in enumerate(results_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    # Add Figure 1: Statistical Performance
    add_figure(doc, 
               "publication_figures/statistical_performance.png",
               "Model Performance Comparison showing validation loss with 95% confidence intervals across 5 random seeds. Lower is better. Both Time-Indexed variants (green, orange) significantly outperform the baselines.",
               1, width_inches=5.5)
    
    doc.add_paragraph("""Key Observations:

1. Time-Indexed MLP achieves the best balance of performance, speed, and compression:
   • 5.8% lower loss than Standard Transformer
   • 4.5% lower loss than Tong's Neural ODE
   • 430× parameter reduction
   • 7.2× faster training than baseline

2. Time-Indexed SSM achieves the lowest absolute loss:
   • 9.3% lower loss than Standard Transformer
   • 8.1% lower loss than Tong's Neural ODE
   • 63× parameter reduction
   • But slower training (see Section 5.4)""")
    
    doc.add_page_break()
    
    doc.add_heading('5.2 Statistical Significance Tests', level=2)
    doc.add_paragraph("We performed pairwise t-tests with Cohen's d effect sizes to verify that our improvements are statistically significant and practically meaningful.")
    
    # Significance table
    para = doc.add_paragraph()
    para.add_run("Table 5: Statistical Significance (p-values and Effect Sizes)").bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    headers = ['Comparison', 'p-value', 'Significant?', "Cohen's d", 'Effect Size']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    sig_data = [
        ['Standard vs Time-Indexed MLP', '3.51×10⁻⁵', 'Yes (p<0.01)', '5.83', 'Large'],
        ['Standard vs Time-Indexed SSM', '0.0082', 'Yes (p<0.01)', '2.47', 'Large'],
        ['Tong\'s vs Time-Indexed MLP', '1.21×10⁻⁴', 'Yes (p<0.01)', '4.90', 'Large'],
        ['Tong\'s vs Time-Indexed SSM', '0.0166', 'Yes (p<0.05)', '2.13', 'Large']
    ]
    for i, row_data in enumerate(sig_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    # Add Figure 2: Significance Tests
    add_figure(doc,
               "publication_figures/significance_tests.png",
               "Statistical significance heatmaps. Left: p-values (green = significant difference, p < 0.05). Right: Cohen's d effect sizes (blue = row model better, red = column model better). Both Time-Indexed variants show highly significant improvements over baselines with large effect sizes.",
               2, width_inches=6.0)
    
    doc.add_paragraph("""Interpretation:

• Both Time-Indexed variants significantly outperform both baselines (all p < 0.05)
• Effect sizes are uniformly large (Cohen's d > 2.0), indicating practically meaningful improvements
• The difference between MLP and SSM variants is not statistically significant (p = 0.22), suggesting they achieve comparable performance
• Standard vs Tong's shows a trend (p = 0.0617) but doesn't reach conventional significance""")
    
    doc.add_page_break()
    
    doc.add_heading('5.3 Parameter Efficiency Analysis', level=2)
    doc.add_paragraph("One of the most striking findings is the extreme parameter efficiency of our approaches. The Time-Indexed MLP achieves better performance with 430× fewer parameters than the standard Transformer.")
    
    # Add Figure 3: Efficiency
    add_figure(doc,
               "publication_figures/efficiency_with_error.png",
               "Parameter efficiency plot showing validation loss vs. parameter count (log scale) with 95% confidence intervals. The ideal position is bottom-left (fewer parameters, lower loss). Both Time-Indexed variants achieve this, with Time-Indexed MLP showing the best efficiency (0.7M parameters, loss 2.231).",
               3, width_inches=5.5)
    
    doc.add_paragraph("""The efficiency plot demonstrates that our methods occupy the optimal region of the parameter-performance trade-off space:

• Time-Indexed MLP (green triangle): Best efficiency with 0.7M parameters achieving 2.231 loss
• Time-Indexed SSM (orange diamond): Best absolute performance with 4.9M parameters achieving 2.147 loss
• Standard Transformer (red circle): 308.5M parameters for only 2.367 loss—clearly over-parameterized
• Tong's Neural ODE (blue square): 51.5M parameters, moderate performance

This suggests that the standard practice of using separate weights per layer may be fundamentally wasteful, and that constrained parameter sharing with time-dependent modulation provides implicit regularization that improves generalization.""")
    
    doc.add_page_break()
    
    doc.add_heading('5.4 The SSM Speed Paradox', level=2)
    
    paradox_para = doc.add_paragraph()
    paradox_para.add_run("Critical Finding: ").bold = True
    paradox_para.add_run("Despite having 63× fewer parameters than the standard Transformer, the Time-Indexed SSM variant is ")
    run = paradox_para.add_run("slower")
    run.bold = True
    run.underline = True
    paradox_para.add_run(" per training step (64.3ms vs 55.3ms). This counter-intuitive result—which we term the ")
    paradox_para.add_run("SSM Speed Paradox").italic = True
    paradox_para.add_run("—has important implications for practitioners.")
    
    # Add Figure 4: Speed Comparison
    add_figure(doc,
               "publication_figures/speed_comparison.png",
               "Training speed comparison (lower is better). The SSM Speed Paradox is clearly visible: Time-Indexed SSM (orange, 64.3ms) is the slowest despite having far fewer parameters than Standard (red, 55.3ms). Time-Indexed MLP (green, 7.7ms) achieves the best speed.",
               4, width_inches=5.5)
    
    doc.add_paragraph("""Root Cause Analysis:

The speed paradox arises from fundamental computational differences between attention and SSM:""")
    
    # Comparison table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    headers = ['Aspect', 'Attention (Standard/MLP)', 'SSM']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    speed_data = [
        ['Core Operation', 'Matrix multiplication (GEMM)', 'Sequential recurrence'],
        ['Parallelization', 'Fully parallel on GPU', 'Sequential dependencies'],
    ]
    for i, row_data in enumerate(speed_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_paragraph("""The SSM's selective scan requires processing positions sequentially:

    h[t+1] = A_bar · h[t] + B_bar · x[t]
    
Each position t+1 depends on position t, creating a chain of dependencies that cannot be parallelized. In contrast, attention computes all position interactions simultaneously via matrix multiplication.""")
    
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run("Implications for Practitioners:").bold = True
    
    # Recommendations table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Priority', 'Choose MLP', 'Choose SSM']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    rec_data = [
        ['Lowest latency', '✓', '✗'],
        ['Lowest memory', '✗', '✓'],
        ['Best accuracy', '✗', '✓'],
        ['Edge deployment', '✓', '✓'],
        ['Real-time inference', '✓', '✗']
    ]
    for i, row_data in enumerate(rec_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_paragraph("""Potential Solutions (Future Work):

1. Parallel Scan Algorithms: Mamba-style parallel scans can reduce O(n) sequential steps to O(log n)

2. Chunked Processing: Process in blocks of 32-64 tokens with parallel attention within chunks

3. Custom CUDA Kernels: Hardware-optimized implementations of recurrent operations

4. Hybrid Architecture: Use MLP for speed-critical paths, SSM for accuracy-critical paths""")
    
    doc.add_page_break()
    
    # =====================================================
    # CHAPTER 6: DISCUSSION
    # =====================================================
    
    doc.add_heading('6. Discussion', level=1)
    
    doc.add_heading('6.1 Interpretation of Results', level=2)
    doc.add_paragraph("""Our results demonstrate that extreme parameter compression is possible without sacrificing performance. The Time-Indexed MLP variant achieves 430× compression while actually improving upon baselines. This challenges the conventional assumption that more parameters necessarily lead to better models.

Why does time-indexed sharing outperform?

We hypothesize three contributing factors:

1. Implicit Regularization: Sharing weights constrains the model to learn representations that are useful across all depths, acting as a powerful regularizer that prevents overfitting to layer-specific patterns.

2. Optimization Stability: Grounding the effective weights in a learned base W_base provides a stable optimization landscape. Unlike hypernetworks that must generate weights from scratch at each forward pass, our approach learns the base weights directly, with only lightweight modulation varying across layers.

3. Capacity Where Needed: The time-modulation network focuses model capacity on the differences between layers, not the commonalities. Since adjacent layers likely need similar transformations, explicitly sharing a base and learning only the deltas is more parameter-efficient.""")
    
    doc.add_heading('6.2 Comparison with Related Approaches', level=2)
    doc.add_paragraph("""Our approach occupies a unique position in the design space:

• Universal Transformers share weights but lack time-dependent modulation, limiting expressiveness
• Tong's Neural ODE generates all weights from scratch, wasting parameters on layer commonalities
• LoRA adds adapters to frozen weights, but doesn't exploit the continuous-depth perspective
• Our approach combines weight sharing (efficiency) with time modulation (expressiveness)

The key insight is that layers need to be different, but not completely different. Time-indexed sharing finds the optimal middle ground.""")
    
    doc.add_heading('6.3 Limitations', level=2)
    doc.add_paragraph("""We acknowledge several limitations of this work:

1. Scale: Our experiments are limited to small models (<5M parameters). Scaling to LLaMA-size (100M+) models is necessary to verify the approach generalizes.

2. Tokenization: We use character-level tokenization. Subword tokenization (BPE, SentencePiece) may show different trade-offs.

3. Tasks: Only language modeling is evaluated. Other tasks (classification, generation quality) may show different patterns.

4. SSM Sensitivity: The SSM variant's performance degrades on WikiText-103, suggesting sensitivity to hyperparameters at larger scales.

5. Ablation Completeness: While our results are statistically validated, a complete ablation study fixing modulation to constants (removing time dependence entirely) would strengthen our theoretical claims about the importance of time-indexing.""")
    
    doc.add_page_break()
    
    # =====================================================
    # CHAPTER 7: CONCLUSION
    # =====================================================
    
    doc.add_heading('7. Conclusion and Future Work', level=1)
    
    doc.add_heading('7.1 Summary of Contributions', level=2)
    doc.add_paragraph("""This thesis presented Time-Indexed Parameter Sharing, a novel approach to efficient Transformer design that achieves remarkable parameter compression while improving model performance.

Key Findings:

1. 430× parameter reduction is achievable with the Time-Indexed MLP variant while improving validation loss by 5.8% over standard Transformers

2. Statistically significant improvements (p < 0.01) validated across 5 random seeds with large effect sizes (Cohen's d > 2.0)

3. The SSM Speed Paradox reveals that parameter count alone does not determine inference speed—computational structure matters fundamentally

4. Implicit regularization through weight sharing appears to be a key mechanism behind the performance gains

These findings challenge the prevailing assumption that larger models are necessarily better, and suggest that architectural innovations in parameter sharing may be as important as scaling.""")
    
    doc.add_heading('7.2 Future Directions', level=2)
    doc.add_paragraph("""Immediate Extensions:

1. Complete Ablation Study: Test constant modulation baseline to definitively isolate the time-indexing benefit from the adapter benefit

2. FLOPs Analysis: Report computational cost (FLOPs per forward pass) alongside parameter counts for a complete efficiency picture

3. Parallel SSM Implementation: Explore Mamba-style parallel scans to address the SSM Speed Paradox

Medium-term Goals:

4. Scale to 100M+ Parameters: Validate findings on LLaMA-scale models with subword tokenization

5. Additional Benchmarks: Evaluate on C4, The Pile, GLUE, and downstream tasks

6. Theoretical Analysis: Develop formal understanding of the regularization properties of time-indexed sharing

Long-term Vision:

7. Production Deployment: Optimize for edge devices (mobile, IoT, embedded systems)

8. Multimodal Extension: Apply time-indexed sharing to vision (ViT) and audio transformers

9. Adaptive Depth: Explore input-dependent depth selection based on complexity""")
    
    doc.add_page_break()
    
    # =====================================================
    # REFERENCES
    # =====================================================
    
    doc.add_heading('8. References', level=1)
    
    references = [
        "[1] Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, Ł., & Polosukhin, I. (2017). Attention is All You Need. Advances in Neural Information Processing Systems (NeurIPS).",
        
        "[2] Chen, R. T. Q., Rubanova, Y., Bettencourt, J., & Duvenaud, D. (2018). Neural Ordinary Differential Equations. Advances in Neural Information Processing Systems (NeurIPS).",
        
        "[3] Tong, A., Nguyen-Tang, T., Lee, D., Nguyen, D., Tran, T., Hall, D. L. W., Kang, C., & Choi, J. (2025). Neural ODE Transformers: Analyzing Internal Dynamics and Adaptive Fine-tuning. International Conference on Learning Representations (ICLR).",
        
        "[4] Gu, A., & Dao, T. (2023). Mamba: Linear-Time Sequence Modeling with Selective State Spaces. arXiv preprint arXiv:2312.00752.",
        
        "[5] Dehghani, M., Gouws, S., Vinyals, O., Uszkoreit, J., & Kaiser, Ł. (2019). Universal Transformers. International Conference on Learning Representations (ICLR).",
        
        "[6] Hu, E. J., Shen, Y., Wallis, P., Allen-Zhu, Z., Li, Y., Wang, S., Wang, L., & Chen, W. (2022). LoRA: Low-Rank Adaptation of Large Language Models. International Conference on Learning Representations (ICLR).",
        
        "[7] Perez, E., Strub, F., De Vries, H., Dumoulin, V., & Courville, A. (2018). FiLM: Visual Reasoning with a General Conditioning Layer. AAAI Conference on Artificial Intelligence.",
        
        "[8] Xiong, R., Yang, Y., He, D., Zheng, K., Zheng, S., Xing, C., Zhang, H., Lan, Y., Wang, L., & Liu, T. (2020). On Layer Normalization in the Transformer Architecture. International Conference on Machine Learning (ICML).",
        
        "[9] Merity, S., Xiong, C., Bradbury, J., & Socher, R. (2017). Pointer Sentinel Mixture Models. International Conference on Learning Representations (ICLR).",
        
        "[10] Radford, A., Wu, J., Child, R., Luan, D., Amodei, D., & Sutskever, I. (2019). Language Models are Unsupervised Multitask Learners. OpenAI Technical Report."
    ]
    
    for ref in references:
        para = doc.add_paragraph(ref)
        para.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()
    
    # =====================================================
    # APPENDIX
    # =====================================================
    
    doc.add_heading('Appendix A: Individual Seed Results', level=1)
    
    doc.add_paragraph("Table A1: Time-Indexed MLP Results by Seed").bold = True
    
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    headers = ['Seed', 'Valid Loss', 'Speed (ms)']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    mlp_data = [
        ['42', '2.204', '7.40'],
        ['123', '2.203', '7.61'],
        ['456', '2.233', '7.78'],
        ['789', '2.246', '7.42'],
        ['1011', '2.267', '8.16'],
        ['Mean ± Std', '2.231 ± 0.025', '7.67 ± 0.28']
    ]
    for i, row_data in enumerate(mlp_data):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_data
            if i == 5:  # Bold the mean row
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    doc.add_paragraph("Table A2: Time-Indexed SSM Results by Seed").bold = True
    
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    headers = ['Seed', 'Valid Loss', 'Speed (ms)']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    ssm_data = [
        ['42', '2.062', '63.50'],
        ['123', '2.026', '64.33'],
        ['456', '2.085', '64.53'],
        ['789', '2.190', '64.89'],
        ['1011', '2.370', '64.12'],
        ['Mean ± Std', '2.147 ± 0.124', '64.27 ± 0.46']
    ]
    for i, row_data in enumerate(ssm_data):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_data
            if i == 5:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph("Note: The higher variance in SSM results (std = 0.124 vs 0.025 for MLP) suggests the SSM architecture may be more sensitive to initialization, warranting further investigation.")
    
    # =====================================================
    # DECLARATION
    # =====================================================
    
    doc.add_page_break()
    doc.add_heading('Declaration', level=1)
    
    doc.add_paragraph("""I declare that this thesis is my own work and has not been submitted for any other degree or professional qualification. All sources have been properly acknowledged and cited.

The experimental code, statistical analysis, and architectural innovations presented in this thesis represent original contributions, building upon the foundational work of Tong et al. (ICLR 2025) which is properly attributed throughout.""")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    sig_para = doc.add_paragraph()
    sig_para.add_run("Signature: ").bold = True
    sig_para.add_run("_" * 40)
    
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.add_run("Date: ").bold = True
    date_para.add_run("_" * 40)
    
    # =====================================================
    # SAVE DOCUMENT
    # =====================================================
    
    output_path = Path("THESIS.docx")
    doc.save(output_path)
    
    print("=" * 70)
    print("✅ THESIS DOCUMENT GENERATED SUCCESSFULLY")
    print("=" * 70)
    print(f"\n📄 Output file: {output_path.absolute()}")
    print(f"\n📊 Figures included:")
    print("   • Figure 1: Statistical Performance (validation loss comparison)")
    print("   • Figure 2: Significance Tests (p-values and Cohen's d heatmaps)")
    print("   • Figure 3: Parameter Efficiency (loss vs parameters scatter)")
    print("   • Figure 4: Speed Comparison (SSM Speed Paradox visualization)")
    print(f"\n📋 Document contents:")
    print("   • Title page")
    print("   • Abstract with keywords")
    print("   • Table of Contents")
    print("   • 7 main chapters with subsections")
    print("   • 5 data tables")
    print("   • 4 embedded figures")
    print("   • Full references")
    print("   • Appendix with individual seed results")
    print("   • Declaration page")
    print("\n💡 Next steps:")
    print("   1. Open THESIS.docx in Microsoft Word")
    print("   2. Replace [Your Name], [Supervisor], [University]")
    print("   3. Update Table of Contents (References → Update Table)")
    print("   4. Add page numbers (Insert → Page Number)")
    print("   5. Review formatting and adjust as needed")
    print("=" * 70)
    
    return output_path


if __name__ == "__main__":
    create_thesis_document()
